from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from docx.shared import Pt
from model import *

from controller.api import api
from view.view_upload import view_upload
from view.view_edit import view_edit
from view.view_manage import view_manage
from view.view_record import view_record
import threading
from docx.enum.text import WD_ALIGN_PARAGRAPH
from asr_server import start

app = Flask(__name__)
app.register_blueprint(api)
app.register_blueprint(view_upload)
app.register_blueprint(view_record)
app.register_blueprint(view_edit)
app.register_blueprint(view_manage)
# app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config["SECRET_KEY"] = b'_5#y2L"F4Q8z\n\xec]/'
app.config['MONGODB_SETTINGS'] = {
    'host': 'mongodb://localhost/detect_audio'
}
db.init_app(app)


@app.route('/', methods=['GET', 'POST'])
def homepage():
    id_audio = request.args.get('id_audio')
    if id_audio == None:
        return render_template('home.html', file_audio="", lyrics=[], spell_mistake=[])
    else:
        audio = FileAudio.objects.get(id_audio=id_audio)
        return render_template('home.html', file_audio = audio.path_file, lyrics = audio.lyrics\
            ,spell_mistake=audio.spell_mistake,  id_audio=audio.id_audio)

@app.route('/tab', methods=['GET', "POST"])
def tabpage():
    id_audio = request.args.get('id_audio')
    if id_audio == None:
        return render_template('tab.html', file_audio="", lyrics=[], spell_mistake=[])
    else:
        audio = FileAudio.objects.get(id_audio=id_audio)
        return render_template('tab.html', file_audio = audio.path_file, lyrics = audio.lyrics\
            ,spell_mistake=audio.spell_mistake,  id_audio=audio.id_audio)


@app.route('/tab_rt', methods=['GET', "POST"])
def tab_rt_page():
    id_audio = request.args.get('id_audio')
    # start()
    # t1 = threading.Thread(name="loop", target=start).start()
    # t1.daemon = True
    # t1.start()
    # if id_audio == None:
    return render_template('tab_rt.html', file_audio="", lyrics=[], spell_mistake=[])
    # else:
    #     audio = FileAudio.objects.get(id_audio=id_audio)
    #     return render_template('index.html', file_audio = audio.path_file, lyrics = audio.lyrics\
    #         ,spell_mistake=audio.spell_mistake,  id_audio=audio.id_audio)
 


@app.route('/edit', methods=['GET'])
def editpage():
    id_audio = request.args.get('id_audio')
    audio = FileAudio.objects.get(id_audio=id_audio)
    return render_template('edit.html', file_audio= audio.path_file, lyrics = audio.lyrics\
        ,spell_mistake=audio.spell_mistake,  id_audio=audio.id_audio)

@app.route('/export-files/<id_audio>')
def return_files(id_audio):
    data = FileAudio.objects(id_audio=id_audio).first()
    res = []
    file_name = data['path_file'].split('/')[-1]
    for line in data['lyrics']:
        res.append(line.split(']')[-1])
    path_txt = data['path_file'][:-3] + 'txt'
    with open(path_txt, 'w', encoding='utf-8') as output:
        output.write(', '.join(res))
    return send_file(path_txt, attachment_filename= file_name[:-3] + 'txt', as_attachment=True)

@app.route('/export-audio/<id_audio>')
def return_audio(id_audio):
    data = FileAudio.objects(id_audio=id_audio).first()
    file_name = data['path_file'].split('/')[-1]
    return send_file(data['path_file'], attachment_filename= file_name, as_attachment=True)

@app.route('/export-doc/<id_audio>')
def return_files_docx(id_audio):
    # https://192.168.1.4:8080/export-doc?classify=upload
    res = []
    data = FileAudio.objects(id_audio=id_audio).first()
    file_name = data['path_file'].split('/')[-1]
    for line in data['lyrics']:
        res.append(line.split(']')[-1])
    path_txt = data['path_file'][:-3] + 'docx'
    document = Document()
    # document.add_heading(file_name, 0).add_run()
    myfile = ', '.join(res)
    # myfile = re.sub(r'[^\x00-\x7F]+|\x0c',' ', myfile)
    # print(myfile)
    run = document.add_paragraph()
    run.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    run=run.add_run(myfile)
    font = run.font
#    font.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font.name = 'Times New Roman'
    font.size = Pt(14)
    # document.add_paragraph(myfile)
    print(path_txt)
    document.save(path_txt)
    
    return send_file(path_txt, attachment_filename= file_name[:-3] + 'docx', as_attachment=True)


if __name__ == '__main__':
    # from waitress import serve
    from argparse import ArgumentParser
    parser = ArgumentParser()
    parser.add_argument('-p', '--port', default=5001,
                        type=int, help='port to listen on')
    parser.add_argument('--host', default='0.0.0.0',
                        type=str, help='port to listen on')
    args = parser.parse_args()
    host = args.host
    port = args.port
    app.run(host=host, port=port, debug=True, threaded=True ,ssl_context=('cert.pem', 'key.pem'))
    # serve(app=app, host=host, port=port,ssl_context=('cert.pem', 'key.pem'))
#    start()

